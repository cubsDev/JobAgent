import os
from flask import Flask, render_template, request, jsonify, redirect, url_for, send_file
from config import Config
from models import db, SearchConfig, JobPost
import math
from docx import Document
from markdown import markdown
from bs4 import BeautifulSoup

def markdown_to_docx(md_text, output_path):
    # Very basic naive markdown to docx converter
    document = Document()
    html = markdown(md_text)
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup:
        if element.name == 'h1':
            document.add_heading(element.text, level=1)
        elif element.name == 'h2':
            document.add_heading(element.text, level=2)
        elif element.name == 'h3':
            document.add_heading(element.text, level=3)
        elif element.name == 'p':
            document.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                document.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                document.add_paragraph(li.text, style='List Number')

    document.save(output_path)
    return output_path

def create_app(config_class=Config):
    app = Flask(__name__)
    app.config.from_object(config_class)

    db.init_app(app)

    with app.app_context():
        # Create database tables if they don't exist
        db.create_all()
        # Seed default SearchConfig if empty
        if not SearchConfig.query.first():
            default_config = SearchConfig(search_term="Software Engineer", location="Remote", results_wanted=30, hours_old=72, score_threshold=7)
            db.session.add(default_config)
            db.session.commit()

    @app.route('/')
    def index():
        from sqlalchemy import case, or_
        page = request.args.get('page', 1, type=int)
        status_filter = request.args.get('status', 'all')
        
        query = JobPost.query
        
        if status_filter != 'all':
            query = query.filter_by(status=status_filter)
            
        config = SearchConfig.query.first()
        active_locations = []
        if config and config.location:
            active_locations = [loc.strip().split(',')[0].lower() for loc in config.location.split('|') if loc.strip()]
            
        if active_locations:
            conditions = [db.func.lower(JobPost.location).contains(city) for city in active_locations]
            is_active_loc = case(
                (or_(*conditions), 1), 
                else_=0
            )
            
            # Sort by active location first, then relevance score
            jobs_pagination = query.order_by(
                is_active_loc.desc(),
                JobPost.relevance_score.desc()
            ).paginate(page=page, per_page=12, error_out=False)
        else:
            # Sort by relevance score descending
            jobs_pagination = query.order_by(JobPost.relevance_score.desc()).paginate(page=page, per_page=12, error_out=False)
        
        return render_template('index.html', jobs=jobs_pagination.items, pagination=jobs_pagination, status_filter=status_filter)

    @app.route('/job/<int:id>')
    def job_detail(id):
        job = JobPost.query.get_or_404(id)
        return render_template('job_detail.html', job=job)

    @app.route('/job/<int:id>/tailor', methods=['POST'])
    def tailor_job_cv(id):
        from agents.tailor import tailor_cv
        import os
        
        job = JobPost.query.get_or_404(id)
        # Load master_cv
        master_cv_path = os.path.join(app.root_path, 'master_cv.txt')
        with open(master_cv_path, 'r', encoding='utf-8') as f:
            master_cv_content = f.read()
            
        # Run tailor
        tailored_md = tailor_cv(master_cv=master_cv_content, job_description=job.description, job_title=job.title)
        job.tailored_cv_md = tailored_md
        db.session.commit()
        
        return redirect(url_for('job_detail', id=job.id))

    @app.route('/job/<int:id>/status', methods=['POST'])
    def update_job_status(id):
        job = JobPost.query.get_or_404(id)
        data = request.get_json()
        new_status = data.get('status')
        if new_status in ['new', 'reviewed', 'applied', 'rejected']:
            job.status = new_status
            db.session.commit()
            return jsonify({"success": True, "status": new_status})
        return jsonify({"success": False, "error": "Invalid status"}), 400

    @app.route('/job/<int:id>/export', methods=['POST'])
    def export_job_cv(id):
        job = JobPost.query.get_or_404(id)
        if not job.tailored_cv_md:
            return "CV not tailored yet", 400
            
        outputs_dir = os.path.join(app.root_path, 'static', 'outputs')
        os.makedirs(outputs_dir, exist_ok=True)
        docx_filename = f"Tailored_CV_{job.id}.docx"
        docx_path = os.path.join(outputs_dir, docx_filename)
        
        markdown_to_docx(job.tailored_cv_md, docx_path)
        job.docx_path = docx_filename
        db.session.commit()
        
        return send_file(docx_path, as_attachment=True, download_name=docx_filename)

    @app.route('/settings', methods=['GET', 'POST'])
    def settings():
        config = SearchConfig.query.first()
        master_cv_path = os.path.join(app.root_path, 'master_cv.txt')
        
        if request.method == 'POST':
            # Update Config
            config.search_term = request.form.get('search_term')
            config.location = '|'.join(request.form.getlist('location'))
            config.results_wanted = int(request.form.get('results_wanted', 30))
            config.hours_old = int(request.form.get('hours_old', 72))
            config.score_threshold = int(request.form.get('score_threshold', 7))
            
            config.scoring_engine = request.form.get('scoring_engine', 'ollama')
            config.keywords = request.form.get('keywords', '')
            
            db.session.commit()
            
            # Update Master CV
            new_cv_content = request.form.get('master_cv_content')
            if new_cv_content is not None:
                with open(master_cv_path, 'w', encoding='utf-8') as f:
                    f.write(new_cv_content)
                    
            return redirect(url_for('settings'))
            
        # GET Load Master CV
        cv_content = ""
        if os.path.exists(master_cv_path):
            with open(master_cv_path, 'r', encoding='utf-8') as f:
                cv_content = f.read()
                
        return render_template('settings.html', config=config, master_cv=cv_content)

    @app.route('/run-now', methods=['POST'])
    def run_pipeline_now():
        from pipeline import run_pipeline, pipeline_state
        import threading
        if pipeline_state.get('is_running', False):
            return jsonify({"status": "Already running"}), 400
            
        # Run asynchronously so we don't block the request
        thread = threading.Thread(target=run_pipeline, args=(app,))
        thread.start()
        return jsonify({"status": "Pipeline started"})

    @app.route('/api/pipeline_status', methods=['GET'])
    def get_pipeline_status():
        from pipeline import pipeline_state
        return jsonify(pipeline_state)

    @app.route('/api/jobs', methods=['GET'])
    def api_jobs():
        jobs = JobPost.query.all()
        return jsonify([
            {
                "id": j.id,
                "title": j.title,
                "company": j.company,
                "score": j.relevance_score,
                "status": j.status,
                "url": j.url
            } for j in jobs
        ])
        
    @app.route('/api/ollama_status', methods=['GET'])
    def ollama_status():
        import requests
        try:
            ollama_url = app.config.get('OLLAMA_BASE_URL', 'http://localhost:11434')
            res = requests.get(ollama_url, timeout=2)
            # Ollama root endpoint returns "Ollama is running"
            is_up = res.status_code == 200
            return jsonify({"status": "up" if is_up else "down"})
        except:
            return jsonify({"status": "down"})

    return app

if __name__ == '__main__':
    app = create_app()
    app.run(debug=True)
